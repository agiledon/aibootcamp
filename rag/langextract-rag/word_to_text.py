"""
Word文档转文本工具

这个模块提供了将Word文档转换为文本字符串的功能。
支持处理单个文件或整个文件夹中的Word文档。
"""

import os
import glob
from pathlib import Path
from typing import List, Optional, Union
import logging

# 尝试导入不同的Word文档处理库
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("Warning: python-docx not available. Install with: pip install python-docx")

try:
    import mammoth
    MAMMOTH_AVAILABLE = True
except ImportError:
    MAMMOTH_AVAILABLE = False
    print("Warning: mammoth not available. Install with: pip install mammoth")

# 配置日志
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


def extract_text_from_docx_python_docx(file_path: str) -> str:
    """
    使用python-docx库提取Word文档文本
    
    Args:
        file_path: Word文档的路径
        
    Returns:
        提取的文本字符串
    """
    try:
        doc = Document(file_path)
        text_parts = []
        
        # 提取段落文本
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text.strip())
        
        # 提取表格文本
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_parts.append(" | ".join(row_text))
        
        return "\n".join(text_parts)
    
    except Exception as e:
        logger.error(f"Error extracting text from {file_path} using python-docx: {e}")
        return ""


def extract_text_from_docx_mammoth(file_path: str) -> str:
    """
    使用mammoth库提取Word文档文本
    
    Args:
        file_path: Word文档的路径
        
    Returns:
        提取的文本字符串
    """
    try:
        with open(file_path, "rb") as docx_file:
            result = mammoth.extract_raw_text(docx_file)
            return result.value
    except Exception as e:
        logger.error(f"Error extracting text from {file_path} using mammoth: {e}")
        return ""


def extract_text_from_word_document(file_path: str, method: str = "auto") -> str:
    """
    从Word文档中提取文本
    
    Args:
        file_path: Word文档的路径
        method: 提取方法 ("auto", "python-docx", "mammoth")
        
    Returns:
        提取的文本字符串
    """
    if not os.path.exists(file_path):
        logger.error(f"File not found: {file_path}")
        return ""
    
    file_path = os.path.abspath(file_path)
    
    if method == "auto":
        # 自动选择可用的最佳方法
        if DOCX_AVAILABLE:
            text = extract_text_from_docx_python_docx(file_path)
            if text.strip():
                return text
        
        if MAMMOTH_AVAILABLE:
            text = extract_text_from_docx_mammoth(file_path)
            if text.strip():
                return text
        
        logger.error("No suitable library available for extracting text from Word documents")
        return ""
    
    elif method == "python-docx" and DOCX_AVAILABLE:
        return extract_text_from_docx_python_docx(file_path)
    
    elif method == "mammoth" and MAMMOTH_AVAILABLE:
        return extract_text_from_docx_mammoth(file_path)
    
    else:
        logger.error(f"Method '{method}' not available or not supported")
        return ""


def convert_documents_to_text(
    folder_path: str,
    file_pattern: str = "*.docx",
    method: str = "auto",
    combine_all: bool = True,
    separator: str = "\n\n" + "="*50 + "\n\n"
) -> Union[str, List[str]]:
    """
    将指定文件夹下的Word文档转换为文本
    
    Args:
        folder_path: 文档文件夹路径
        file_pattern: 文件匹配模式，默认为"*.docx"
        method: 提取方法 ("auto", "python-docx", "mammoth")
        combine_all: 是否将所有文档合并为一个字符串
        separator: 文档之间的分隔符（仅在combine_all=True时使用）
        
    Returns:
        如果combine_all=True，返回合并的文本字符串
        如果combine_all=False，返回文本字符串列表
    """
    folder_path = os.path.abspath(folder_path)
    
    if not os.path.exists(folder_path):
        logger.error(f"Folder not found: {folder_path}")
        return "" if combine_all else []
    
    # 查找匹配的文件
    pattern = os.path.join(folder_path, file_pattern)
    matching_files = glob.glob(pattern)
    
    if not matching_files:
        logger.warning(f"No files matching pattern '{file_pattern}' found in {folder_path}")
        return "" if combine_all else []
    
    logger.info(f"Found {len(matching_files)} files to process")
    
    # 处理每个文件
    texts = []
    for file_path in sorted(matching_files):
        logger.info(f"Processing: {os.path.basename(file_path)}")
        text = extract_text_from_word_document(file_path, method)
        
        if text.strip():
            # 添加文件名作为标题
            filename = os.path.basename(file_path)
            text_with_title = f"# {filename}\n\n{text}"
            texts.append(text_with_title)
        else:
            logger.warning(f"No text extracted from {file_path}")
    
    if combine_all:
        return separator.join(texts)
    else:
        return texts


def save_text_to_file(text: str, output_path: str) -> bool:
    """
    将文本保存到文件
    
    Args:
        text: 要保存的文本
        output_path: 输出文件路径
        
    Returns:
        是否保存成功
    """
    try:
        # 确保输出目录存在
        output_dir = os.path.dirname(output_path)
        if output_dir:  # 只有当目录不为空时才创建
            os.makedirs(output_dir, exist_ok=True)
        
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(text)
        logger.info(f"Text saved to: {output_path}")
        return True
    except Exception as e:
        logger.error(f"Error saving text to {output_path}: {e}")
        return False


def main():
    """示例用法"""
    # 示例：处理document文件夹下的所有Word文档
    document_folder = "document"
    
    if os.path.exists(document_folder):
        print("Converting Word documents to text...")
        
        # 转换为合并的文本字符串
        combined_text = convert_documents_to_text(
            folder_path=document_folder,
            file_pattern="*.docx",
            method="auto",
            combine_all=True
        )
        
        if combined_text:
            # 保存到文件
            output_file = "extracted_text.txt"
            if save_text_to_file(combined_text, output_file):
                print(f"Successfully extracted {len(combined_text)} characters of text")
                print(f"Text saved to: {output_file}")
            else:
                print("Failed to save text to file")
        else:
            print("No text was extracted")
    else:
        print(f"Document folder not found: {document_folder}")


if __name__ == "__main__":
    main()
